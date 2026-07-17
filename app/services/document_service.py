"""文档解析与分块服务 —— 支持 PDF / TXT / MD / DOCX / CSV / HTML"""
import csv
import io
import os
import re
from typing import List

from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader


class DocumentChunk:
    """文档分块"""

    def __init__(self, content: str, metadata: dict):
        self.content = content
        self.metadata = metadata


# 支持的文件扩展名
SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".csv", ".html", ".htm"}


class DocumentService:
    """文档服务：多格式文档解析 + 文本分块"""

    # 网格搜索最优参数（在 24914 字学术论文上评估得出）
    CHUNK_SIZE = 1000
    CHUNK_OVERLAP = 50

    def __init__(self, upload_dir: str = "uploads"):
        self.upload_dir = upload_dir
        os.makedirs(upload_dir, exist_ok=True)

    # ─── 各格式解析器 ─────────────────────────────────

    def _parse_pdf(self, path: str) -> str:
        reader = PdfReader(path)
        parts = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
        return "\n".join(parts)

    def _parse_txt(self, path: str) -> str:
        # 尝试多种编码
        for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                with open(path, encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        return ""

    def _parse_docx(self, path: str) -> str:
        try:
            from docx import Document
            doc = Document(path)
            parts = []
            # 段落文字
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    parts.append(p.text.strip())
            # 表格文字
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)
            return "\n".join(parts)
        except Exception as e:
            raise ValueError(f"Word 文档解析失败: {e}")

    def _parse_csv(self, path: str) -> str:
        rows = []
        for enc in ("utf-8", "gbk", "gb2312"):
            try:
                with open(path, encoding=enc, newline="") as f:
                    reader = csv.reader(f)
                    for r in reader:
                        rows.append(" | ".join(r))
                break
            except (UnicodeDecodeError, csv.Error):
                continue
        return "\n".join(rows)

    def _parse_html(self, path: str) -> str:
        from bs4 import BeautifulSoup
        for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                with open(path, encoding=enc) as f:
                    soup = BeautifulSoup(f.read(), "lxml")
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
        # 移除 script / style 标签
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)

    # ─── 统一入口 ─────────────────────────────────────

    def parse_file(self, file_path: str, filename: str) -> str:
        """根据扩展名自动选择解析器，返回纯文本"""
        ext = os.path.splitext(filename)[1].lower()
        parsers = {
            ".pdf": self._parse_pdf,
            ".txt": self._parse_txt,
            ".md": self._parse_txt,
            ".docx": self._parse_docx,
            ".csv": self._parse_csv,
            ".html": self._parse_html,
            ".htm": self._parse_html,
        }
        parser = parsers.get(ext)
        if parser is None:
            raise ValueError(f"不支持的文件格式: {ext}")
        try:
            return parser(file_path)
        except ValueError:
            raise  # 重新抛出已包装的错误
        except Exception as e:
            raise ValueError(f"文档解析失败（{ext}）: {e}")

    # ─── 文本分块 ─────────────────────────────────────

    # 中文法律条文编号正则：第X条、第X章、第X节
    _ARTICLE_RE = re.compile(
        r'(?:^|\n)\s*(第[一二三四五六七八九十百千\d]+条)\s*',
        re.MULTILINE,
    )
    _SECTION_RE = re.compile(
        r'(?:^|\n)\s*(第[一二三四五六七八九十\d]+[章节编])\s*[^\n]*',
        re.MULTILINE,
    )

    # 最小块长度：低于此值的短法条会与相邻条合并
    MIN_CHUNK_CHARS = 150

    def chunk_text(self, text: str, filename: str, source_type: str = "uploaded") -> List["DocumentChunk"]:
        """按法条号分块：每个 chunk 是一条或多条完整法条

        对于法律文档（含「第X条」标记），按法条边界切分；
        对于普通文档（无「第X条」），回退到 RecursiveCharacterTextSplitter。
        """
        # 检测是否为法律文档（含法条编号）
        if not self._ARTICLE_RE.search(text):
            return self._chunk_by_char(text, filename, source_type)

        return self._chunk_by_article(text, filename, source_type)

    def _chunk_by_article(self, text: str, filename: str, source_type: str) -> List["DocumentChunk"]:
        """按法条号分块：每条/每组法条一个 chunk"""
        # 提取所有章节标题位置
        sections = [(m.start(), m.group().strip()) for m in self._SECTION_RE.finditer(text)]

        # 按「第X条」切分
        parts = self._ARTICLE_RE.split(text)
        # parts[0] = 前言，parts[1] = "第一条", parts[2] = 第一条内容，
        #   parts[3] = "第二条", parts[4] = 第二条内容, ...

        if len(parts) < 3:
            # 只有一条法条或切分失败，回退
            return self._chunk_by_char(text, filename, source_type)

        # 收集 (法条号, 法条内容)
        articles = []
        preamble = parts[0].strip()
        if preamble:
            articles.append(("前言", preamble))

        for i in range(1, len(parts) - 1, 2):
            label = parts[i].strip()
            content = parts[i + 1].strip() if i + 1 < len(parts) else ""
            if content:
                articles.append((label, content))

        # 查找每个法条所属的章/节
        def find_section(pos):
            sec = ""
            for start, title in sections:
                if start < pos:
                    sec = title
            return sec

        # 合并短法条 + 添加章节上下文 → 生成 chunk
        chunks = []
        buffer_label = ""
        buffer_text = ""
        buffer_pos = 0
        pos_so_far = len(preamble) if preamble else 0

        for label, content in articles:
            pos_so_far = text.find(label, pos_so_far)
            if pos_so_far < 0:
                pos_so_far = 0

            if buffer_text and len(buffer_text) + len(content) < self.MIN_CHUNK_CHARS * 2:
                # 合并到当前 buffer
                buffer_text += "\n" + label + " " + content
                continue
            elif buffer_text and len(buffer_text) < self.MIN_CHUNK_CHARS:
                # Buffer 太短，继续合并
                buffer_text += "\n" + label + " " + content
                continue

            # 输出当前 buffer
            if buffer_text:
                section = find_section(buffer_pos)
                header = f"{section}\n{buffer_label}" if section else buffer_label
                chunks.append(DocumentChunk(
                    content=header + "\n" + buffer_text if section else buffer_text,
                    metadata={"source": filename, "chunk_index": len(chunks), "source_type": source_type},
                ))

            buffer_label = label
            buffer_text = content
            buffer_pos = pos_so_far

        # 最后一个 buffer
        if buffer_text:
            section = find_section(buffer_pos)
            header = f"{section}\n{buffer_label}" if section else buffer_label
            chunks.append(DocumentChunk(
                content=header + "\n" + buffer_text if section else buffer_text,
                metadata={"source": filename, "chunk_index": len(chunks), "source_type": source_type},
            ))

        return chunks

    def _chunk_by_char(self, text: str, filename: str, source_type: str) -> List["DocumentChunk"]:
        """字符级分块（非法律文档回退方案）"""
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.CHUNK_SIZE,
            chunk_overlap=self.CHUNK_OVERLAP,
            separators=["\n\n", "\n", "。", "！", "？", "；", " ", ""],
            length_function=len,
        )
        chunks = splitter.split_text(text)
        return [
            DocumentChunk(
                content=c,
                metadata={
                    "source": filename,
                    "chunk_index": i,
                    "source_type": source_type,
                },
            )
            for i, c in enumerate(chunks)
        ]

    # ─── 一站式处理 ───────────────────────────────────

    def process_file(self, file_path: str, filename: str, source_type: str = "uploaded") -> List[DocumentChunk]:
        """解析 + 分块，一站式"""
        text = self.parse_file(file_path, filename)
        if not text.strip():
            return []
        return self.chunk_text(text, filename, source_type=source_type)

    # ─── 兼容旧方法名 ─────────────────────────────────

    def parse_pdf(self, path: str) -> str:
        return self._parse_pdf(path)

    def process_pdf(self, file_path: str, filename: str) -> List[DocumentChunk]:
        return self.process_file(file_path, filename)
